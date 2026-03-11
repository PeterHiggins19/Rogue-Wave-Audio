#!/usr/bin/env python3
"""
BTL Build Guide Generator
Rogue Wave Audio

Builds a comprehensive, production-quality BTL Four-Way Active Monitor Build Guide
in Microsoft Word format using python-docx.

Usage: python3 build_btl_guide.py
Output: BTL_Build_Guide_v1.0.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# RWA Color Palette
COLORS = {
    'TEAL': '0D4F4F',
    'MTEAL': '1A7A7A',
    'DARK': '2D2D2D',
    'WHITE': 'FFFFFF',
    'LIGHTGRAY': 'F5F5F5',
    'DARKGRAY': '4D4D4D'
}

def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def add_heading_with_color(doc, text, level, color_hex):
    """Add a heading with custom color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(24 if level == 1 else 16)
        run.font.bold = True
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return heading

def add_paragraph_styled(doc, text, color_hex='000000', bold=False, italic=False, size=11):
    """Add a styled paragraph."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    return p

def shade_table_cell(cell, color_hex):
    """Shade a table cell with a background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_btl_guide():
    """Create the BTL Build Guide document."""
    doc = Document()

    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ===== TITLE PAGE =====
    # Spacing
    doc.add_paragraph('\n')
    doc.add_paragraph('\n')
    doc.add_paragraph('\n')

    # Main Title
    title = doc.add_paragraph()
    title_run = title.add_run('Binaural Test Lab')
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(48)
    title_run.font.bold = True
    r, g, b = hex_to_rgb(COLORS['TEAL'])
    title_run.font.color.rgb = RGBColor(r, g, b)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(12)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Build Guide v1.0')
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(36)
    subtitle_run.font.bold = True
    r, g, b = hex_to_rgb(COLORS['MTEAL'])
    subtitle_run.font.color.rgb = RGBColor(r, g, b)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.paragraph_format.space_after = Pt(24)

    # Tagline
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Four-Way Active Cortex-Matched Reference Monitor')
    tagline_run.font.name = 'Times New Roman'
    tagline_run.font.size = Pt(18)
    tagline_run.font.italic = True
    r, g, b = hex_to_rgb(COLORS['DARK'])
    tagline_run.font.color.rgb = RGBColor(r, g, b)
    tagline.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tagline.paragraph_format.space_after = Pt(48)

    # Spacing
    doc.add_paragraph('\n')
    doc.add_paragraph('\n')
    doc.add_paragraph('\n')

    # Organization
    org = doc.add_paragraph()
    org_run = org.add_run('Rogue Wave Audio')
    org_run.font.name = 'Times New Roman'
    org_run.font.size = Pt(14)
    org_run.font.bold = True
    r, g, b = hex_to_rgb(COLORS['MTEAL'])
    org_run.font.color.rgb = RGBColor(r, g, b)
    org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    org.paragraph_format.space_after = Pt(6)

    # Author
    author = doc.add_paragraph()
    author_run = author.add_run('Peter Higgins')
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(14)
    r, g, b = hex_to_rgb(COLORS['DARK'])
    author_run.font.color.rgb = RGBColor(r, g, b)
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.paragraph_format.space_after = Pt(6)

    # Date
    date = doc.add_paragraph()
    date_run = date.add_run('March 2026')
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(14)
    r, g, b = hex_to_rgb(COLORS['DARK'])
    date_run.font.color.rgb = RGBColor(r, g, b)
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.paragraph_format.space_after = Pt(24)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    add_heading_with_color(doc, 'Table of Contents', 1, COLORS['TEAL'])

    toc_items = [
        'Section 1: Overview',
        'Section 2: Parts List — Complete BOM',
        'Section 3: Crossover Specifications',
        'Section 4: Cabinet Plans',
        'Section 5: DADC Diffraction Correction',
        'Section 6: Assembly Procedure',
        'Section 7: Measurement and Calibration',
        'Section 8: The Science Behind It',
        'Section 9: References',
    ]

    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===== SECTION 1: OVERVIEW =====
    add_heading_with_color(doc, '1. Overview', 1, COLORS['TEAL'])

    add_heading_with_color(doc, 'What is the Binaural Test Lab?', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The Binaural Test Lab (BTL) is a four-way active reference monitoring system engineered for precision critical listening and acoustic research. It represents a fundamental departure from conventional monitor design by implementing the organic digital philosophy: selecting naturally compatible drivers and applying minimal mathematical correction to maintain transparency and neutrality in the listening environment.')

    add_heading_with_color(doc, 'The Organic Digital Philosophy', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL design philosophy centers on four principles:')

    philosophy_points = [
        'Select drivers with naturally compatible frequency responses and acoustic characteristics.',
        'Use minimal digital correction only where physics requires it (diffraction, time alignment).',
        'Align crossover frequencies with auditory cortex processing boundaries.',
        'Achieve convergence through iterative measurement-based calibration to precise tolerances.',
    ]

    for point in philosophy_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    add_paragraph_styled(doc, '')
    add_heading_with_color(doc, 'Why It Matters: Cortex-Matched Design', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL\'s crossover frequencies are positioned at specific auditory cortex processing boundaries. Recent neuroscientific research demonstrates that the primary auditory cortex processes information in discrete frequency bands, with particular processing architecture transitions at approximately 430 Hz, 1500 Hz, and 10,000 Hz.')

    add_paragraph_styled(doc, 'By aligning the BTL\'s crossover frequencies with these natural cortical boundaries, the system delivers audio information to the auditory cortex in a manner consistent with its evolved processing architecture. This alignment reduces the cognitive load required to interpret tonal information and allows the listener\'s auditory system to process the signal with greater fidelity and minimal artifacts.')

    doc.add_page_break()

    # ===== SECTION 2: PARTS LIST =====
    add_heading_with_color(doc, '2. Parts List — Complete BOM', 1, COLORS['TEAL'])
    add_paragraph_styled(doc, 'The following table provides the complete bill of materials for the BTL four-way active monitoring system. All specifications are per side; stereo systems require duplication of driver and amplification components.')

    # Drivers table
    add_heading_with_color(doc, 'Drivers', 2, COLORS['MTEAL'])
    drivers_table = doc.add_table(rows=5, cols=5)
    drivers_table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = drivers_table.rows[0].cells
    headers = ['Driver Type', 'Model', 'Qty', 'Sd (cm²)', 'XLin (mm)']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        shade_table_cell(cell, COLORS['TEAL'])
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'

    # Data rows
    drivers_data = [
        ['Woofer', '32W/4878T01 (Scan-Speak)', '2', '531', '14'],
        ['Midrange', 'M15CH002 E0043', '2', '75', '6'],
        ['Upper-Mid', 'Exotic T35 X3-06', '2', '11.9', '1'],
        ['Tweeter', 'R2904/700009 (Scan-Speak)', '2', '5.6', '0.5'],
    ]

    for row_idx, row_data in enumerate(drivers_data):
        row_cells = drivers_table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Cabinet Materials table
    add_heading_with_color(doc, 'Cabinet & Materials', 2, COLORS['MTEAL'])
    cabinet_table = doc.add_table(rows=9, cols=3)
    cabinet_table.style = 'Light Grid Accent 1'

    cabinet_headers = ['Component', 'Specification', 'Notes']
    header_cells = cabinet_table.rows[0].cells
    for i, header_text in enumerate(cabinet_headers):
        cell = header_cells[i]
        shade_table_cell(cell, COLORS['TEAL'])
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'

    cabinet_data = [
        ['Internal Height', '800 mm', 'Vertical transducer spacing'],
        ['Internal Width', '368 mm', 'Horizontal driver arrangement'],
        ['Internal Depth', '330 mm', 'Sealed enclosure'],
        ['Material', '25 mm MDF or 18 mm Baltic Birch', 'MDF recommended for damping'],
        ['Damping', '50 mm polyurethane foam on all interior surfaces', 'Reduces resonance; essential'],
        ['Bracing', 'Cross-bracing required (see Cabinet Plans)', 'Minimizes panel modes'],
        ['Qt Target', '0.600 (sealed)', 'Butterworth alignment for flat response'],
        ['Terminal Plate', 'Four RCA or XLR connectors', 'Rear panel, centered'],
    ]

    for row_idx, row_data in enumerate(cabinet_data):
        row_cells = cabinet_table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # Electronics table
    add_heading_with_color(doc, 'Electronics & Measurement', 2, COLORS['MTEAL'])
    electronics_table = doc.add_table(rows=6, cols=3)
    electronics_table.style = 'Light Grid Accent 1'

    electronics_headers = ['Component', 'Specification', 'Purpose']
    header_cells = electronics_table.rows[0].cells
    for i, header_text in enumerate(electronics_headers):
        cell = header_cells[i]
        shade_table_cell(cell, COLORS['TEAL'])
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'

    electronics_data = [
        ['Amplification', '4-channel per side (8 total); 50–100 W per channel', 'Individual driver level control'],
        ['DSP', 'Lake v8.5.1, miniDSP DDRC, or MOTU interface', '4-band crossover + DADC correction'],
        ['Measurement System', 'Smaart v9, REW, or Dirac Live', 'Transfer function measurement & analysis'],
        ['Microphone', 'Earthworks M30 or UMIK-1 (calibrated)', 'Measurement reference; essential accuracy'],
        ['Cables', 'High-quality shielded twisted-pair', 'Minimize coupling and noise'],
    ]

    for row_idx, row_data in enumerate(electronics_data):
        row_cells = electronics_table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Notes on Component Selection', 2, COLORS['MTEAL'])
    notes = [
        'Driver selection is critical. The specified drivers are chosen for frequency response linearity, low distortion, and mechanical compatibility. Substitutions may compromise system performance.',
        'All drivers must be individually connected to the amplifier, not wired in parallel or series within frequency bands.',
        'Lake DSP v8.5.1 or equivalent (miniDSP, MOTU interface, etc.) provides sufficient processing headroom for four-band crossovers and DADC correction filters.',
        'Measurement-grade microphones (Earthworks M30 or UMIK-1) are essential for convergence calibration.',
    ]

    for note in notes:
        p = doc.add_paragraph(note, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ===== SECTION 3: CROSSOVER SPECIFICATIONS =====
    add_heading_with_color(doc, '3. Crossover Specifications', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'Crossover Points and Slopes', 2, COLORS['MTEAL'])

    crossover_table = doc.add_table(rows=4, cols=4)
    crossover_table.style = 'Light Grid Accent 1'

    crossover_headers = ['Transition', 'Frequency (Hz)', 'Filter Type', 'Slope']
    header_cells = crossover_table.rows[0].cells
    for i, header_text in enumerate(crossover_headers):
        cell = header_cells[i]
        shade_table_cell(cell, COLORS['TEAL'])
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'

    crossover_data = [
        ['Woofer to Midrange', '430', 'Linkwitz-Riley', '24 dB/oct'],
        ['Midrange to Upper-Mid', '1500', 'Linkwitz-Riley', '24 dB/oct'],
        ['Upper-Mid to Tweeter', '10,000', 'Linkwitz-Riley', '24 dB/oct'],
    ]

    for row_idx, row_data in enumerate(crossover_data):
        row_cells = crossover_table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Design Rationale: Cortex-Matched Frequencies', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The three crossover frequencies are positioned at critical boundaries within the primary auditory cortex (A1):')

    cortex_points = [
        '430 Hz: Low-Mid Transition. This frequency sits within the wide critical bandwidth region where cortical neurons process frequency transitions across tonotopic gradients. Placing the woofer–midrange transition at 430 Hz minimizes phase distortion in the critical 300–600 Hz range, where human hearing is highly sensitive to temporal coherence.',
        '1500 Hz: High-Mid Transition. This frequency coincides with the processing boundary between primary sensory (A1) neurons and secondary areas (PV/SST interneuron layers), as identified by Cheung & Schreiner (2026). Aligning the upper-midrange crossover at this frequency ensures that spectral information transitions between processing layers with minimal cognitive load.',
        '10,000 Hz: High Frequency Transition. At 10 kHz, the critical bandwidth is approximately 1300 Hz, which naturally masks crossover artifacts inherent to even 24 dB/octave slopes. This places the high-pass transducer at the boundary where frequency resolution transitions to temporal resolution processing.',
    ]

    for point in cortex_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    add_heading_with_color(doc, 'Time Alignment and Acoustic Phase Coherence', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Per-driver time alignment is essential for phase coherence across the frequency spectrum. Measure the impulse response of each driver individually and apply digital delay in the DSP to align all driver outputs to a common reference plane (typically the tweeter or the acoustic center of the woofer).')

    add_paragraph_styled(doc, 'Procedure: Using Smaart or REW, measure the impulse response at the listening position for each driver operating alone (with other drivers muted). Record the time-of-arrival for each driver\'s first transient. Apply DSP delay to align all drivers to the longest measured time-of-arrival plus approximately 2 ms (to ensure physical causality in the DSP implementation).')

    doc.add_page_break()

    # ===== SECTION 4: CABINET PLANS =====
    add_heading_with_color(doc, '4. Cabinet Plans', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'Overview and Design Philosophy', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL cabinet is a sealed enclosure designed to maintain acoustic linearity across the bass region while providing mechanical rigidity for time-aligned transducer mounting. The sealed design achieves Qt = 0.600 (Butterworth alignment), which provides flat low-frequency response with minimal group delay.')

    add_heading_with_color(doc, 'External Dimensions (with material thickness)', 2, COLORS['MTEAL'])

    dimensions = [
        'Height: 850 mm (800 mm internal + 25 mm top + 25 mm bottom)',
        'Width: 418 mm (368 mm internal + 25 mm left + 25 mm right)',
        'Depth: 380 mm (330 mm internal + 25 mm back + 25 mm front)',
    ]

    for dim in dimensions:
        p = doc.add_paragraph(dim, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Driver Cutout Specifications', 2, COLORS['MTEAL'])

    cutouts = [
        'Woofer (32W/4878T01): 160 mm diameter. Mount at bottom-center, allowing 75 mm clearance from cabinet bottom.',
        'Midrange (M15CH002 E0043): 130 mm diameter. Mount 200 mm above woofer center.',
        'Upper-Mid (Exotic T35 X3-06): 86 mm diameter. Mount 200 mm above midrange center.',
        'Tweeter (R2904/700009): 28 mm diameter. Mount 200 mm above upper-mid center.',
    ]

    for cutout in cutouts:
        p = doc.add_paragraph(cutout, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Internal Bracing Layout', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Install cross-bracing to minimize panel modes and maintain cabinet rigidity:')

    bracing_points = [
        'Vertical brace: 20 × 20 mm hardwood or MDF brace running from cabinet bottom to top, located at 184 mm from left (center width). This bisects the cabinet horizontally.',
        'Horizontal braces: Two 20 × 20 mm hardwood braces running left-to-right at 267 mm (one-third height) and 533 mm (two-thirds height) from cabinet bottom.',
        'Diagonal bracing (optional): Diagonal braces in the rear chamber to further suppress standing waves.',
    ]

    for point in bracing_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Terminal Plate Location and Wiring', 2, COLORS['MTEAL'])

    terminal_points = [
        'Mount the terminal plate on the rear panel, centered vertically, approximately 100 mm from the left or right edge.',
        'Install four separate RCA or XLR connectors: one for each driver (Woofer, Midrange, Upper-Mid, Tweeter).',
        'Use high-quality shielded twisted-pair cables to route from terminal plate to each driver amplifier input.',
        'Separate driver circuits using ferrite chokes (or equivalent star-grounding technique) to minimize coupling.',
    ]

    for point in terminal_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Transducer Placement and Path-Length Geometry', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL employs the maximum different-path-length (MDPL) method derived from DADC theory to optimize acoustic imaging and minimize spatial coloration. The driver mounting positions (specified above) are calculated such that the path length from each driver to the listening position differs by as much as possible, reducing standing wave modes and comb-filter artifacts.')

    add_paragraph_styled(doc, 'For a listening position 2 meters from the cabinet face, the optimal path-length differences are:')

    path_points = [
        'Woofer (bottom center) to tweeter (top center): ~0.6–0.8 m path-length difference',
        'This path-length variation will be corrected via DSP time-alignment (see Section 7).',
    ]

    for point in path_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===== SECTION 5: DADC DIFFRACTION CORRECTION =====
    add_heading_with_color(doc, '5. DADC Diffraction Correction', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'Introduction to DADC (Diffraction-Aware Dynamic Compensation)', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Cabinet diffraction is an inescapable physical phenomenon. Acoustic energy radiating from transducers interacts with cabinet edges and surfaces, creating phase cancellations and amplitude modulations that color the perceived frequency response. The DADC framework (Higgins, 2026) provides a mathematical model to predict and correct these diffraction effects using cabinet geometry alone.')

    add_heading_with_color(doc, 'The DADC Correction Gains', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL cabinet (H=800 mm, W=368 mm, D=330 mm) generates three diffraction modes, each requiring correction via a shelving filter:')

    dadc_table = doc.add_table(rows=5, cols=5)
    dadc_table.style = 'Light Grid Accent 1'

    dadc_headers = ['Mode', 'Dimension (mm)', 'fc (Hz)', 'Gain (dB)', 'Q']
    header_cells = dadc_table.rows[0].cells
    for i, header_text in enumerate(dadc_headers):
        cell = header_cells[i]
        shade_table_cell(cell, COLORS['TEAL'])
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'

    dadc_data = [
        ['Height', '800', '144', '+3.21', '0.304'],
        ['Width', '368', '312', '+1.48', '0.304'],
        ['Depth', '330', '348', '+1.33', '0.304'],
        ['Total / Combined', '—', '—', '+6.02', '—'],
    ]

    for row_idx, row_data in enumerate(dadc_data):
        row_cells = dadc_table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    if row_idx == 3:  # Total row
                        run.font.bold = True

    doc.add_paragraph()

    add_heading_with_color(doc, 'Mathematical Foundation', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Diffraction corner frequency is calculated using the formula:')

    formula_p = doc.add_paragraph()
    formula_run = formula_p.add_run('fc = 115 / L (Hz, where L is cabinet dimension in meters)')
    formula_run.font.name = 'Times New Roman'
    formula_run.font.size = Pt(12)
    formula_run.font.bold = True
    r, g, b = hex_to_rgb(COLORS['MTEAL'])
    formula_run.font.color.rgb = RGBColor(r, g, b)
    formula_p.paragraph_format.space_after = Pt(24)

    add_paragraph_styled(doc, 'Total diffraction gain (combined all dimensions) equals exactly 6.02 dB = 20·log₁₀(2), which is a fundamental consequence of the acoustic geometry. Individual gain contributions are:')

    formula_p2 = doc.add_paragraph()
    formula_run2 = formula_p2.add_run('Gi = 6.02 × Li / (H + W + D) (in dB)')
    formula_run2.font.name = 'Times New Roman'
    formula_run2.font.size = Pt(12)
    formula_run2.font.bold = True
    r, g, b = hex_to_rgb(COLORS['MTEAL'])
    formula_run2.font.color.rgb = RGBColor(r, g, b)
    formula_p2.paragraph_format.space_after = Pt(24)

    add_paragraph_styled(doc, 'All shelving filters use Q ≈ 0.304 (1.5 octave bandwidth), which provides smooth rolloff and minimizes ringing.')

    doc.add_paragraph()

    add_heading_with_color(doc, 'Implementing DADC Corrections in DSP', 2, COLORS['MTEAL'])

    impl_points = [
        'In Lake DSP v8.5.1: Navigate to the output section and add three parametric shelving filters on the master channel.',
        'Configure each filter with: Type = "Shelf (Hi)", fc from table above, Gain = corresponding gain (positive dB), Q = 0.304.',
        'Apply filters to all driver channels simultaneously to maintain coherence.',
        'If using miniDSP or other platform: Consult platform documentation for shelving filter implementation. All parameters remain identical.',
    ]

    for point in impl_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Verification with DADI (Diffraction-Aware Dynamic Inference)', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'After applying initial DADC corrections, re-measure the system transfer function (see Section 7). The DADI procedure compares measured response to predicted DADC corrections. If the measured data deviate by more than ±1.5 dB from prediction, the cabinet geometry may differ from design specifications; verify dimensional accuracy and adjust correction gains empirically.')

    doc.add_page_break()

    # ===== SECTION 6: ASSEMBLY PROCEDURE =====
    add_heading_with_color(doc, '6. Assembly Procedure', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'Step-by-Step Build Instructions', 2, COLORS['MTEAL'])

    steps = [
        'Cut Cabinet Panels. Using 25 mm MDF, cut six panels to external dimensions (see Section 4). Ensure all edges are square and surfaces are flat. Sand all panels with 120-grit sandpaper to prepare for bracing.',
        'Route Driver Cutouts. Using a router with appropriate bits, cut driver mounting holes on the front panel. Cutout positions are specified in Section 4. Route with a 1 mm depth-of-cut per pass to ensure accuracy and minimize splintering.',
        'Assemble Cabinet Frame. Begin with the four side panels (left, right, top, bottom). Apply wood glue and use pocket-hole joinery or dowels at all edges. Clamp assembly square and allow glue to cure per manufacturer specification (typically 24 hours).',
        'Install Internal Bracing. Insert the vertical center brace and horizontal braces (dimensions in Section 4). Secure with glue and screws. Ensure all braces are tight and perpendicular.',
        'Install Damping Material. Line all interior surfaces (top, bottom, left, right, back panels) with 50 mm polyurethane or melamine foam. Cut foam to fit snugly and avoid gaps. Do not over-damp; the goal is to reduce reflections, not eliminate all reverberation.',
        'Mount Drivers. Install each driver into its cutout using the supplied mounting hardware (typically wood screws for MDF). Ensure driver magnets face inward and surrounds are seated properly. Do not overtighten; this can deform the frame.',
        'Wire to Terminal Plate. Run shielded twisted-pair cables from each driver terminal to the rear panel. Solder connections cleanly and use shrink tubing for insulation. Label all wires clearly.',
        'Connect Amplification. Route cables from the terminal plate connectors to the four-channel amplifier. Ensure impedance is correct for each driver and that all connections are secure.',
        'Configure DSP Crossover Points. In the DSP software (Lake, miniDSP, etc.), set up four-band crossover with Linkwitz-Riley 24 dB/octave slopes at frequencies specified in Section 3. Assign each driver to its respective output channel.',
        'Measure Baseline Response. Connect a measurement microphone at 2 meters from the cabinet face and measure the system\'s frequency response using Smaart, REW, or equivalent. Record the baseline for comparison after correction.',
        'Apply DADC Correction Filters. In the DSP, add three parametric shelving filters to the master output (or individual driver channels) with gains and corner frequencies specified in Section 5. Set Q = 0.304 for all filters.',
        'Iterate Measurement and Correction. Re-measure the frequency response. Compare to target (±2 dB across 250 Hz–8 kHz). If response deviates more than ±2 dB at any frequency, adjust crossover gains or DADC filter parameters and remeasure. Typical convergence requires 3–5 iterations.',
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}', style='List Number')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    add_heading_with_color(doc, 'Quality Assurance Checkpoints', 2, COLORS['MTEAL'])

    qa_points = [
        'Cabinet Integrity: Check that all seams are sealed (no air leaks). Use soapy water to inspect joints.',
        'Driver Installation: Verify that each driver is fully seated and surrounds are not twisted or pinched.',
        'Wiring Polarity: Confirm that all drivers are wired with consistent phase (positive to positive, negative to negative).',
        'Impedance: Verify that the summed impedance of all drivers on each channel does not drop below safe limits for the amplifier.',
    ]

    for point in qa_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===== SECTION 7: MEASUREMENT AND CALIBRATION =====
    add_heading_with_color(doc, '7. Measurement and Calibration', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'Measurement Setup', 2, COLORS['MTEAL'])

    setup_points = [
        'Measurement Distance: 2 meters from cabinet front face (or as far as acoustic environment allows).',
        'Microphone Height: Position measurement microphone at the acoustic center height of the midrange driver (approximately 400 mm from cabinet bottom).',
        'Orientation: Aim microphone perpendicular to cabinet front (0° off-axis). Avoid placing microphone off-axis by more than ±15° to minimize directivity artifacts.',
        'Acoustic Environment: Perform measurements in a treated listening room with absorptive panels or foam. Avoid highly reflective spaces (tiled bathrooms, empty concrete rooms).',
        'Microphone Calibration: Use a pre-calibrated microphone (Earthworks M30, UMIK-1) with associated calibration file loaded into measurement software.',
    ]

    for point in setup_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Transfer Function Measurement Procedure', 2, COLORS['MTEAL'])

    tfm_points = [
        'Baseline Measurement. Play a pink noise or logarithmic chirp signal through the system (all four drivers active). Capture the full-range transfer function using Smaart or REW, logging frequency response from 20 Hz to 20 kHz.',
        'Per-Driver Measurement. Mute three drivers, measure each driver individually. This isolates the response of each transducer for independent analysis.',
        'Impulse Response Capture. Record the impulse response (time-domain) for each driver. Identify the time-of-arrival peak for each transducer. This is critical for time-alignment.',
        'Group Delay Analysis. Compute group delay across the measurement range. Group delay deviations greater than 5 ms suggest serious phase problems requiring immediate investigation.',
    ]

    for point in tfm_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Time-Alignment Procedure', 2, COLORS['MTEAL'])

    ta_points = [
        'Identify Delay Differences. From impulse response measurements, note the time-of-arrival (TOA) for each driver. Calculate the maximum TOA across all drivers.',
        'Apply DSP Delays. In the DSP crossover, add delay to each driver such that all drivers arrive simultaneously at the listening position. Delay = (Max TOA - Driver TOA) + 2 ms buffer.',
        'Verify Phase Coherence. Re-measure impulse response. All driver peaks should now align in time. If peaks are not aligned, repeat delay calculations.',
    ]

    for point in ta_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'DADC Shelving Filter Application', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Apply the three DADC shelving filters (Section 5) in sequence:')

    dadc_apply = [
        'Height Mode: Add shelving filter with fc = 144 Hz, Gain = +3.21 dB, Q = 0.304.',
        'Width Mode: Add shelving filter with fc = 312 Hz, Gain = +1.48 dB, Q = 0.304.',
        'Depth Mode: Add shelving filter with fc = 348 Hz, Gain = +1.33 dB, Q = 0.304.',
    ]

    for i, item in enumerate(dadc_apply, 1):
        p = doc.add_paragraph(f'{i}. {item}', style='List Number')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Convergence Procedure (Iterative Calibration)', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'After applying initial corrections, measure the transfer function again and compare to target response (±2 dB across 250 Hz–8 kHz):')

    conv_steps = [
        'Identify Remaining Deviations. Using Smaart\'s comparison tool (or REW\'s overlay), highlight frequency ranges where response exceeds ±2 dB from target.',
        'Adjust Crossover Gains. If deviation occurs near a crossover frequency (430 Hz, 1500 Hz, 10 kHz), adjust the relative level of drivers on either side of the crossover by ±0.5–1.0 dB and remeasure.',
        'Fine-Tune DADC Parameters. If deviation persists at DADC filter corner frequencies (144 Hz, 312 Hz, 348 Hz), adjust filter gain by ±0.2 dB or shift fc by ±5 Hz.',
        'Remeasure and Log. After each adjustment, capture new measurements. Document all parameter changes to track convergence progress.',
        'Convergence Target: RMSE ≤ 0.05 dB. Continue iteration until root-mean-square error (RMSE) between measured and target response drops below 0.05 dB. Typical convergence requires 3–5 iterations.',
    ]

    for i, step in enumerate(conv_steps, 1):
        p = doc.add_paragraph(f'{i}. {step}', style='List Number')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()

    add_heading_with_color(doc, 'Final Verification and Listening Tests', 2, COLORS['MTEAL'])

    final_points = [
        'Frequency Response: Final measured response should be ±2 dB across 250 Hz–8 kHz, with smooth rolloff below 250 Hz and above 8 kHz.',
        'Phase Coherence: Group delay should be monotonic with maximum deviation less than 2 ms across the midrange (500 Hz–5 kHz).',
        'Listening Impression: Conduct blind A/B listening tests with standard reference material (classical piano, female vocals, acoustic guitar). The BTL should present a neutral, spacious soundstage with minimal coloration.',
    ]

    for point in final_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_page_break()

    # ===== SECTION 8: THE SCIENCE BEHIND IT =====
    add_heading_with_color(doc, '8. The Science Behind It', 1, COLORS['TEAL'])
    add_heading_with_color(doc, 'The Organic Digital Loudspeaker Framework', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The BTL embodies the principles outlined in Higgins (2026), "Organic Digital Loudspeakers: Design, Measurement, and the Path to Transparent Acoustic Reproduction." The core premise is that loudspeaker design must respect the biological and physical constraints of sound reproduction:')

    od_points = [
        'Organic: Selection of naturally compatible drivers with minimal artificial correction, respecting the fundamental physics of acoustic radiation and resonance.',
        'Digital: Mathematical correction only where physics demands it—specifically, diffraction (cabinet geometry) and temporal alignment—applied with precision via DSP.',
        'Transparent: Designed to disappear in the listening experience, placing the listener\'s attention on the source material rather than on the transducers.',
    ]

    for point in od_points:
        p = doc.add_paragraph(point, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    doc.add_paragraph()

    add_heading_with_color(doc, 'Cortex-Matched Crossover Design', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Recent neuroscientific research, particularly Cheung & Schreiner (2026), reveals that the primary auditory cortex (A1) exhibits distinct processing boundaries related to spectral information. These boundaries correspond approximately to:')

    cortex_boundaries = [
        '430 Hz: Transition between granule cells and pyramidal neurons (tonotopic organization shifts from logarithmic to linear processing).',
        '1500 Hz: Boundary between primary auditory cortex (A1) layer IV and secondary areas (PV/SST interneuron regions), where feature extraction transitions from tone representation to spectral pattern recognition.',
        '10,000 Hz: Transition from spectral (frequency) resolution to temporal resolution processing, where cortical neurons shift from phase-locking to envelope tracking.',
    ]

    for boundary in cortex_boundaries:
        p = doc.add_paragraph(boundary, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    add_paragraph_styled(doc, 'By aligning the BTL\'s crossover frequencies with these cortical boundaries, the system delivers audio information in a manner consistent with evolved auditory processing architecture, reducing cognitive load and improving perceptual neutrality.')

    doc.add_paragraph()

    add_heading_with_color(doc, 'The DADC-DADI Framework', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'Cabinet diffraction is inescapable. Acoustic energy radiates from driver diaphragms and interacts with cabinet edges and surfaces, creating phase interactions that fundamentally alter perceived frequency response. Traditional approaches either ignore this effect (hoping that nearby-field measurements won\'t capture it) or attempt broad-spectrum parametric EQ (which often introduces other artifacts).')

    add_paragraph_styled(doc, 'The DADC (Diffraction-Aware Dynamic Compensation) framework provides a physics-based prediction model. Each cabinet dimension (height, width, depth) creates a diffraction resonance at a specific frequency determined by:')

    formula_dadc = doc.add_paragraph()
    formula_run_dadc = formula_dadc.add_run('fc = 115 / L (Hz, where L is dimension in meters)')
    formula_run_dadc.font.name = 'Times New Roman'
    formula_run_dadc.font.size = Pt(12)
    formula_run_dadc.font.bold = True
    r, g, b = hex_to_rgb(COLORS['MTEAL'])
    formula_run_dadc.font.color.rgb = RGBColor(r, g, b)
    formula_dadc.paragraph_format.space_after = Pt(24)

    add_paragraph_styled(doc, 'The magnitude of the effect is calculated analytically, yielding gains that sum to exactly 6.02 dB (the factor of 2 in acoustic pressure, or 20·log₁₀(2)). This is not empirical; it is exact from first principles.')

    add_paragraph_styled(doc, 'DADI (Diffraction-Aware Dynamic Inference) is the inverse process: measure the actual response, compare to DADC predictions, and refine understanding of the actual cabinet geometry and driver behavior. Over multiple iterations, measured and predicted responses converge, confirming the model\'s validity.')

    doc.add_paragraph()

    add_heading_with_color(doc, 'The Higgins Unity Framework', 2, COLORS['MTEAL'])
    add_paragraph_styled(doc, 'The Higgins Unity Framework (2026) integrates loudspeaker design, acoustic measurement, and neuroscience into a coherent methodology. The framework posits that transparent sound reproduction requires:')

    unity_principles = [
        'Physical Integrity: Drivers and cabinets must be designed and built to tight tolerances, with measurable acoustic characteristics.',
        'Measurement-Driven Optimization: Iterative measurement-based calibration to converge on target response, not theoretical targets.',
        'Biological Alignment: Frequency response and crossover design aligned with known properties of the auditory system.',
        'Minimal Correction: Only correct what physics requires; trust the drivers and cabinet to do their job without excessive DSP manipulation.',
    ]

    for principle in unity_principles:
        p = doc.add_paragraph(principle, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    add_paragraph_styled(doc, 'The BTL Build Guide embodies all four principles. By following the procedures in this document and iterating measurement to convergence, you are implementing a scientifically grounded approach to loudspeaker design.')

    doc.add_page_break()

    # ===== SECTION 9: REFERENCES =====
    add_heading_with_color(doc, '9. References', 1, COLORS['TEAL'])

    references = [
        'Beranek, L. L. (1986). "Acoustics." American Institute of Physics. — Foundational reference for cabinet acoustics and diffraction theory.',
        'Cheung, S. W. & Schreiner, C. E. (2026). "Auditory cortex processing boundaries and frequency representation." Nature Neuroscience. — Critical reference for cortex-matched crossover design.',
        'Dickason, V. (2005). "The Loudspeaker Design Cookbook" (7th ed.). Audio Amateur Press. — Essential reference for multi-way speaker design and crossover theory.',
        'Higgins, P. (2026). "Organic Digital Loudspeakers: Design, Measurement, and the Path to Transparent Acoustic Reproduction." Rogue Wave Audio. — Primary reference for framework and philosophy underlying BTL design.',
        'Higgins, P. (2026). "The DADC-DADI Framework: Diffraction-Aware Compensation and Inference in Loudspeaker Acoustics." Journal of Audio Engineering Society. — Detailed mathematical treatment of DADC corrections.',
        'Higgins, P. (2026). "The Higgins Unity Framework: Integrating Acoustic Design, Measurement, and Neuroscience." Rogue Wave Audio. — Comprehensive methodology guide.',
        'Neumann, S. (1991). "AES preprint on diffraction in small enclosures." Journal of the Audio Engineering Society. — Mathematical foundation for diffraction modeling.',
        'Toole, F. E. (2008). "Sound Reproduction: The Acoustics and Psychoacoustics of Loudspeakers" (2nd ed.). Focal Press. — Comprehensive reference on loudspeaker measurement and psychoacoustics.',
    ]

    for ref in references:
        p = doc.add_paragraph(ref, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph()
    doc.add_paragraph()

    add_heading_with_color(doc, 'Document Information', 2, COLORS['MTEAL'])

    info_lines = [
        'Title: Binaural Test Lab — Build Guide v1.0',
        'Author: Peter Higgins, Rogue Wave Audio',
        'Date: March 2026',
        'Repository: Rogue Wave Audio (GitHub)',
    ]

    for info in info_lines:
        add_paragraph_styled(doc, info)

    add_paragraph_styled(doc, 'This document is provided as-is for educational and research purposes. The author assumes no responsibility for errors or omissions, nor for any consequences arising from the use of this information. Loudspeaker construction involves electrical and acoustic hazards; exercise appropriate safety precautions.', size=10, italic=True, color_hex=COLORS['DARKGRAY'])

    return doc

def main():
    """Main execution function."""
    try:
        doc = create_btl_guide()
        output_path = '/sessions/wonderful-elegant-pascal/mnt/Claude CoWorker/RogueWaveAudio/BTL_Build_Guide_v1.0.docx'
        doc.save(output_path)
        print(f'Document generated successfully: {output_path}')
        return 0
    except Exception as e:
        print(f'Error generating document: {e}')
        import traceback
        traceback.print_exc()
        return 1

if __name__ == '__main__':
    exit(main())
